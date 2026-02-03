#!/usr/bin/env python3
"""
Braille Converter - PDF/Image to Braille (English & Arabic) with Word export
Supports: PDF, PNG, JPG/JPEG → TXT, DOCX (normal), DOCX (braille)
"""

import os
import sys
import re
import argparse
from pathlib import Path
from typing import Optional


class BrailleConverter:
    ENGLISH_BRAILLE = {
        'a': '⠁', 'b': '⠃', 'c': '⠉', 'd': '⠙', 'e': '⠑', 'f': '⠋', 'g': '⠛', 'h': '⠓',
        'i': '⠊', 'j': '⠚', 'k': '⠅', 'l': '⠇', 'm': '⠍', 'n': '⠝', 'o': '⠕', 'p': '⠏',
        'q': '⠟', 'r': '⠗', 's': '⠎', 't': '⠞', 'u': '⠥', 'v': '⠧', 'w': '⠺', 'x': '⠭',
        'y': '⠽', 'z': '⠵', '1': '⠼⠁', '2': '⠼⠃', '3': '⠼⠉', '4': '⠼⠙', '5': '⠼⠑',
        '6': '⠼⠋', '7': '⠼⠛', '8': '⠼⠓', '9': '⠼⠊', '0': '⠼⠚', '.': '⠲', ',': '⠂',
        '?': '⠦', '!': '⠖', ';': '⠆', ':': '⠒', '-': '⠤', '(': '⠐⠣', ')': '⠐⠜',
        '/': '⠸⠌', '"': '⠐⠦', "'": '⠄', '@': '⠈⠁', '#': '⠼', '$': '⠈⠎', '%': '⠨⠴',
        '&': '⠈⠯', '*': '⠐⠔', '+': '⠐⠖', '=': '⠐⠶', '<': '⠐⠣', '>': '⠐⠜',
        '[': '⠈⠣', ']': '⠈⠜', '{': '⠸⠣', '}': '⠸⠜', '\\': '⠸⠡', '|': '⠸⠳',
        '~': '⠸⠔', '_': '⠸⠤', ' ': ' ', '\t': '    ', '\n': '\n', '\r': ''
    }
    
    ARABIC_BRAILLE = {
        'ا': '⠁', 'ب': '⠃', 'ت': '⠞', 'ث': '⠹', 'ج': '⠚', 'ح': '⠱', 'خ': '⠭', 'د': '⠙',
        'ذ': '⠮', 'ر': '⠗', 'ز': '⠵', 'س': '⠎', 'ش': '⠩', 'ص': '⠯', 'ض': '⠫', 'ط': '⠾',
        'ظ': '⠿', 'ع': '⠷', 'غ': '⠣', 'ف': '⠋', 'ق': '⠟', 'ك': '⠅', 'ل': '⠇', 'م': '⠍',
        'ن': '⠝', 'ه': '⠓', 'و': '⠺', 'ي': '⠊', 'ى': '⠊', 'ة': '⠡', 'ء': '⠄', 'ؤ': '⠺',
        'ئ': '⠊', 'أ': '⠁', 'إ': '⠁', 'آ': '⠜', '٠': '⠼⠚', '١': '⠼⠁', '٢': '⠼⠃',
        '٣': '⠼⠉', '٤': '⠼⠙', '٥': '⠼⠑', '٦': '⠼⠋', '٧': '⠼⠛', '٨': '⠼⠓', '٩': '⠼⠊',
        '؟': '⠦', '،': '⠂', '؛': '⠆', '٪': '⠨⠴', '.': '⠲', ',': '⠂', '?': '⠦',
        '!': '⠖', ';': '⠆', ':': '⠒', '-': '⠤', '(': '⠐⠣', ')': '⠐⠜', ' ': ' ',
        '\t': '    ', '\n': '\n', '\r': '', 'َ': '', 'ُ': '', 'ِ': '', 'ّ': '',
        'ْ': '', 'ً': '', 'ٌ': '', 'ٍ': '', 'ـ': ''
    }
    
    CAPITAL_INDICATOR = '⠠'
    
    def __init__(self):
        self.arabic_pattern = re.compile(r'[\u0600-\u06FF\u0750-\u077F]+')
        self.english_pattern = re.compile(r'[a-zA-Z]+')
    
    def is_arabic(self, char):
        return '\u0600' <= char <= '\u06FF' or '\u0750' <= char <= '\u077F'
    
    def detect_language(self, text):
        arabic_count = len(self.arabic_pattern.findall(text))
        english_count = len(self.english_pattern.findall(text))
        if arabic_count > 0 and english_count > 0:
            return 'mixed'
        return 'arabic' if arabic_count > 0 else 'english'
    
    def text_to_braille(self, text):
        if not text:
            return ""
        result = []
        for char in text:
            if self.is_arabic(char):
                result.append(self.ARABIC_BRAILLE.get(char, char))
            else:
                if char.isupper():
                    result.append(self.CAPITAL_INDICATOR)
                    char = char.lower()
                result.append(self.ENGLISH_BRAILLE.get(char, char))
        return ''.join(result)
    
    def get_language_stats(self, text):
        arabic_count = sum(1 for c in text if self.is_arabic(c))
        english_count = sum(1 for c in text if c.isalpha() and not self.is_arabic(c))
        return {
            'arabic_chars': arabic_count,
            'english_chars': english_count,
            'total_chars': len(text),
            'primary_language': self.detect_language(text)
        }


class FileProcessor:
    def __init__(self):
        self.converter = BrailleConverter()
    
    def extract_pdf(self, path):
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("Install pdfplumber: pip install pdfplumber")
        
        extracted = []
        with pdfplumber.open(path) as pdf:
            total = len(pdf.pages)
            print(f"Processing PDF with {total} page(s)...")
            for i, page in enumerate(pdf.pages, 1):
                print(f"  Page {i}/{total}...", end='\r')
                text = page.extract_text()
                if text:
                    extracted.append(text)
                else:
                    print(f"\n  Warning: Page {i} empty or image-only")
            print()
        
        if not extracted:
            raise ValueError("No text extracted from PDF")
        return '\n\n'.join(extracted)
    
    def extract_image(self, path, languages=None):
        try:
            from PIL import Image
            import pytesseract
        except ImportError as e:
            raise ImportError(f"Install pytesseract and Pillow: pip install pytesseract Pillow")
        
        print("Performing OCR...")
        image = Image.open(path)
        
        if languages is None:
            try:
                print("  Trying eng+ara...")
                text = pytesseract.image_to_string(image, lang='eng+ara')
            except:
                print("  Arabic not available, using eng only")
                text = pytesseract.image_to_string(image, lang='eng')
        else:
            text = pytesseract.image_to_string(image, lang=languages)
        
        if not text.strip():
            raise ValueError("No text extracted from image")
        
        lang = self.converter.detect_language(text)
        print(f"  Detected: {lang}")
        return text
    
    def process_file(self, path):
        if not os.path.exists(path):
            raise FileNotFoundError(f"File not found: {path}")
        
        ext = Path(path).suffix.lower()
        if ext == '.pdf':
            return self.extract_pdf(path)
        elif ext in ['.png', '.jpg', '.jpeg']:
            return self.extract_image(path)
        else:
            raise ValueError(f"Unsupported: {ext}. Use .pdf, .png, .jpg, .jpeg")
    
    def save_normal_docx(self, text, output_path):
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from docx.oxml.ns import qn
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")
        
        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(12)
        
        for para_text in text.split('\n'):
            if not para_text.strip():
                doc.add_paragraph()
                continue
            
            lang = self.converter.detect_language(para_text)
            para = doc.add_paragraph()
            run = para.add_run(para_text)
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            
            if lang in ['arabic', 'mixed']:
                para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                pPr = para._element.get_or_add_pPr()
                pPr.set(qn('w:bidi'), '1')
            else:
                para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        doc.save(output_path)
        return output_path
    
    def save_braille_docx(self, text, output_path):
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")
        
        doc = Document()
        doc.styles['Normal'].font.name = 'Arial Unicode MS'
        doc.styles['Normal'].font.size = Pt(14)
        
        for para_text in text.split('\n'):
            if not para_text.strip():
                doc.add_paragraph()
                continue
            
            para = doc.add_paragraph()
            run = para.add_run(para_text)
            run.font.name = 'Arial Unicode MS'
            run.font.size = Pt(14)
            para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        doc.save(output_path)
        return output_path
    
    def convert(self, input_path, output_txt=None):
        print(f"\n{'='*60}")
        print("Braille Converter (English & Arabic)")
        print(f"{'='*60}")
        print(f"Input: {input_path}\n")
        
        # Extract text
        try:
            text = self.process_file(input_path)
            print(f"\n✓ Extracted {len(text)} characters")
        except Exception as e:
            print(f"\n✗ Error: {e}")
            sys.exit(1)
        
        # Language stats
        stats = self.converter.get_language_stats(text)
        print(f"\nLanguage: {stats['primary_language'].upper()}")
        if stats['english_chars'] > 0:
            print(f"  English: {stats['english_chars']}")
        if stats['arabic_chars'] > 0:
            print(f"  Arabic: {stats['arabic_chars']}")
        
        # Convert to Braille
        print("\nConverting to Braille...")
        braille = self.converter.text_to_braille(text)
        
        # Paths
        base = Path(input_path).stem
        txt_path = output_txt or f"{base}_braille.txt"
        normal_docx = f"{base}_normal.docx"
        braille_docx = f"{base}_braille.docx"
        
        # Save TXT
        try:
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(braille)
            print(f"✓ Text: {txt_path}")
        except Exception as e:
            print(f"✗ Text error: {e}")
            sys.exit(1)
        
        # Save Word documents
        print("\nGenerating Word documents...")
        try:
            self.save_normal_docx(text, normal_docx)
            print(f"✓ Normal: {normal_docx}")
        except ImportError as e:
            print(f"⚠ Warning: {e}")
        except Exception as e:
            print(f"⚠ Normal docx error: {e}")
        
        try:
            self.save_braille_docx(braille, braille_docx)
            print(f"✓ Braille: {braille_docx}")
        except ImportError as e:
            print(f"⚠ Warning: {e}")
        except Exception as e:
            print(f"⚠ Braille docx error: {e}")
        
        print(f"\n{'='*60}\n")
        return txt_path


def main():
    parser = argparse.ArgumentParser(
        description='Convert PDF/images to Braille (English & Arabic)',
        epilog="""
Examples:
  %(prog)s document.pdf
  %(prog)s image.png -o output.txt

Outputs:
  <input>_braille.txt   - Braille text
  <input>_normal.docx   - Normal text (RTL for Arabic)
  <input>_braille.docx  - Braille in Word
        """
    )
    parser.add_argument('input_file', help='PDF or image file')
    parser.add_argument('-o', '--output', dest='output_file', 
                       help='Output text file path (optional)', default=None)
    
    args = parser.parse_args()
    processor = FileProcessor()
    processor.convert(args.input_file, args.output_file)


if __name__ == '__main__':
    main()
