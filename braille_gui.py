#!/usr/bin/env python3
"""
Braille Converter GUI — PDF / Image ➜ Braille (English & Arabic)
Exports: TXT (braille), DOCX (normal), DOCX (braille)

Dependencies (install via pip):
    pip install pdfplumber pytesseract Pillow python-docx

Tesseract OCR binary must also be installed on your system.
"""

import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox
from pathlib import Path
import threading
import traceback


class BrailleConverter:
    ENGLISH_BRAILLE = {
        'a': '⠁', 'b': '⠃', 'c': '⠉', 'd': '⠙', 'e': '⠑', 'f': '⠋', 'g': '⠛', 'h': '⠓',
        'i': '⠊', 'j': '⠚', 'k': '⠅', 'l': '⠇', 'm': '⠍', 'n': '⠝', 'o': '⠕', 'p': '⠏',
        'q': '⠟', 'r': '⠗', 's': '⠎', 't': '⠞', 'u': '⠥', 'v': '⠧', 'w': '⠺', 'x': '⠭',
        'y': '⠽', 'z': '⠵', '1': '⠼⠁', '2': '⠼⠃', '3': '⠼⠉', '4': '⠼⠙', '5': '⠼⠑',
        '6': '⠼⠋', '7': '⠼⠛', '8': '⠼⠓', '9': '⠼⠊', '0': '⠼⠚', '.': '⠲', ',': '⠂',
        '?': '⠦', '!': '⠖', ';': '⠆', ':': '⠒', '-': '⠤', '(': '⠐⠣', ')': '⠐⠜',
        '/': '⠸⠌', '"': '⠐⠦', "'": '⠄', '@': '⠈⠁', '#': '⠼', '$': '⠈⠎', '%': '⠨⠴',
        '&': '⠈⠯', '*': '⠐⠔', '+': '⠐⠖', '=': '⠐⠶', '<': '⠐⠣', '>': '⠐⠜',
        '[': '⠈⠣', ']': '⠈⠜', '{': '⠸⠣', '}': '⠸⠜', '\\': '⠸⠡', '|': '⠸⠳',
        '~': '⠸⠔', '_': '⠸⠤', ' ': ' ', '\t': '    ', '\n': '\n', '\r': ''
    }

    ARABIC_BRAILLE = {
        'ا': '⠁', 'ب': '⠃', 'ت': '⠞', 'ث': '⠹', 'ج': '⠚', 'ح': '⠱', 'خ': '⠭', 'د': '⠙',
        'ذ': '⠮', 'ر': '⠗', 'ز': '⠵', 'س': '⠎', 'ش': '⠩', 'ص': '⠯', 'ض': '⠫', 'ط': '⠾',
        'ظ': '⠿', 'ع': '⠷', 'غ': '⠣', 'ف': '⠋', 'ق': '⠟', 'ك': '⠅', 'ل': '⠇', 'م': '⠍',
        'ن': '⠝', 'ه': '⠓', 'و': '⠺', 'ي': '⠊', 'ى': '⠊', 'ة': '⠡', 'ء': '⠄', 'ؤ': '⠺',
        'ئ': '⠊', 'أ': '⠁', 'إ': '⠁', 'آ': '⠜', '٠': '⠼⠚', '١': '⠼⠁', '٢': '⠼⠃',
        '٣': '⠼⠉', '٤': '⠼⠙', '٥': '⠼⠑', '٦': '⠼⠋', '٧': '⠼⠛', '٨': '⠼⠓', '٩': '⠼⠊',
        '؟': '⠦', '،': '⠂', '؛': '⠆', '٪': '⠨⠴', '.': '⠲', ',': '⠂', '?': '⠦',
        '!': '⠖', ';': '⠆', ':': '⠒', '-': '⠤', '(': '⠐⠣', ')': '⠐⠜', ' ': ' ',
        '\t': '    ', '\n': '\n', '\r': '', 'َ': '', 'ُ': '', 'ِ': '', 'ّ': '',
        'ْ': '', 'ً': '', 'ٌ': '', 'ٍ': '', 'ـ': ''
    }

    CAPITAL_INDICATOR = '⠠'

    def __init__(self):
        self.arabic_pattern = re.compile(r'[\u0600-\u06FF\u0750-\u077F]+')
        self.english_pattern = re.compile(r'[a-zA-Z]+')

    def is_arabic(self, char):
        return '\u0600' <= char <= '\u06FF' or '\u0750' <= char <= '\u077F'

    def detect_language(self, text):
        arabic_count  = len(self.arabic_pattern.findall(text))
        english_count = len(self.english_pattern.findall(text))
        if arabic_count > 0 and english_count > 0:
            return 'mixed'
        return 'arabic' if arabic_count > 0 else 'english'

    def text_to_braille(self, text):
        if not text:
            return ""
        result = []
        for char in text:
            if self.is_arabic(char):
                result.append(self.ARABIC_BRAILLE.get(char, char))
            else:
                if char.isupper():
                    result.append(self.CAPITAL_INDICATOR)
                    char = char.lower()
                result.append(self.ENGLISH_BRAILLE.get(char, char))
        return ''.join(result)

    def get_language_stats(self, text):
        arabic_count  = sum(1 for c in text if self.is_arabic(c))
        english_count = sum(1 for c in text if c.isalpha() and not self.is_arabic(c))
        return {
            'arabic_chars':    arabic_count,
            'english_chars':   english_count,
            'total_chars':     len(text),
            'primary_language': self.detect_language(text)
        }


class FileProcessor:

    def __init__(self, log_callback=None):
        self.converter   = BrailleConverter()
        self._log        = log_callback or print

    def extract_pdf(self, path):
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber is required.  Install it:\n  pip install pdfplumber")

        extracted = []
        with pdfplumber.open(path) as pdf:
            total = len(pdf.pages)
            self._log(f"Processing PDF … {total} page(s) found.")
            for i, page in enumerate(pdf.pages, 1):
                self._log(f"  Reading page {i}/{total} …")
                text = page.extract_text()
                if text:
                    extracted.append(text)
                else:
                    self._log(f"  ⚠ Page {i} is empty or image-only.")
        if not extracted:
            raise ValueError("No text could be extracted from this PDF.")
        return '\n\n'.join(extracted)

    @staticmethod
    def _resolve_tesseract():
        import shutil
        if shutil.which("tesseract"):
            return
        candidates = [
            r"C:\Program Files\tesseract-ocr\tesseract.exe",
            r"C:\Program Files (x86)\tesseract-ocr\tesseract.exe",
        ]
        for candidate in candidates:
            if os.path.isfile(candidate):
                import pytesseract
                pytesseract.pytesseract.tesseract_cmd = candidate
                return
        raise FileNotFoundError(
            "Tesseract binary not found.\n\n"
            "Install it from:  https://github.com/UB-Manip/tesseract/releases\n"
            "Then add its folder to PATH, or place it in:\n"
            "  C:\\Program Files\\tesseract-ocr\\"
        )

    def extract_image(self, path, languages=None):
        try:
            from PIL import Image
            import pytesseract
        except ImportError:
            raise ImportError(
                "pytesseract & Pillow are required.\n"
                "  pip install pytesseract Pillow"
            )

        self._resolve_tesseract()

        self._log("Running OCR on image …")
        image = Image.open(path)

        if languages is None:
            try:
                self._log("  Trying eng+ara …")
                text = pytesseract.image_to_string(image, lang='eng+ara')
            except Exception:
                self._log("  Arabic pack unavailable — falling back to eng only.")
                text = pytesseract.image_to_string(image, lang='eng')
        else:
            text = pytesseract.image_to_string(image, lang=languages)

        if not text.strip():
            raise ValueError("No text could be extracted from this image.")
        self._log(f"  Detected language: {self.converter.detect_language(text).upper()}")
        return text

    def process_file(self, path):
        if not os.path.exists(path):
            raise FileNotFoundError(f"File not found:\n  {path}")
        ext = Path(path).suffix.lower()
        if ext == '.pdf':
            return self.extract_pdf(path)
        elif ext in ('.png', '.jpg', '.jpeg'):
            return self.extract_image(path)
        else:
            raise ValueError(f"Unsupported file type: '{ext}'.\nSupported: .pdf  .png  .jpg  .jpeg")

    def save_normal_docx(self, text, output_path):
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from docx.oxml.ns import qn
        except ImportError:
            raise ImportError("python-docx is required.  Install it:\n  pip install python-docx")

        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(12)

        for para_text in text.split('\n'):
            if not para_text.strip():
                doc.add_paragraph()
                continue
            lang  = self.converter.detect_language(para_text)
            para  = doc.add_paragraph()
            run   = para.add_run(para_text)
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            if lang in ('arabic', 'mixed'):
                para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                pPr = para._element.get_or_add_pPr()
                pPr.set(qn('w:bidi'), '1')
            else:
                para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        doc.save(output_path)
        return output_path

    def save_braille_docx(self, text, output_path):
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            raise ImportError("python-docx is required.  Install it:\n  pip install python-docx")

        doc = Document()
        doc.styles['Normal'].font.name = 'Arial Unicode MS'
        doc.styles['Normal'].font.size = Pt(14)

        for para_text in text.split('\n'):
            if not para_text.strip():
                doc.add_paragraph()
                continue
            para = doc.add_paragraph()
            run  = para.add_run(para_text)
            run.font.name = 'Arial Unicode MS'
            run.font.size = Pt(14)
            para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        doc.save(output_path)
        return output_path


COLORS = {
    "bg":           "#1e1e2e",
    "surface":      "#2a2a3d",
    "surface_alt":  "#33334a",
    "accent":       "#7c6aff",
    "accent_hover": "#9585ff",
    "accent_press": "#6350e0",
    "text":         "#e2e2ee",
    "text_dim":     "#9090a8",
    "text_dark":    "#1e1e2e",
    "success":      "#50fa7b",
    "warning":      "#f1fa8c",
    "error":        "#ff5555",
    "border":       "#44445a",
    "log_bg":       "#141420",
    "log_success":  "#50fa7b",
    "log_warning":  "#f1fa8c",
    "log_error":    "#ff5555",
    "log_info":     "#8be9fd",
}

FONT_FAMILY  = "Segoe UI"
FONT_BODY    = (FONT_FAMILY, 11)
FONT_LOG     = ("Consolas", 10)
FONT_STAT    = (FONT_FAMILY, 10, "bold")


class ToolTip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text   = text
        self.tip    = None
        widget.bind("<Enter>", self._show)
        widget.bind("<Leave>", self._hide)

    def _show(self, event=None):
        self.tip = tk.Toplevel(self.widget)
        self.tip.wm_overrideredirect(True)
        self.tip.wm_geometry(f"+{self.widget.winfo_rootx()+15}+{self.widget.winfo_rooty()-10}")
        lbl = tk.Label(self.tip, text=self.text, background="#2a2a3d", foreground="#e2e2ee",
                       font=(FONT_FAMILY, 9), relief="flat", padx=6, pady=3, bd=0)
        lbl.pack()

    def _hide(self, event=None):
        if self.tip:
            self.tip.destroy()
            self.tip = None


class StyledButton(tk.Canvas):

    def __init__(self, master, text, command=None, accent=True, width=160, height=38, **kw):
        try:
            parent_bg = master.cget("bg")
        except (tk.TclError, AttributeError):
            parent_bg = COLORS["bg"]
        super().__init__(master, width=int(width), height=int(height),
                         bg=parent_bg, highlightthickness=0)
        self._text    = text
        self._cmd     = command
        self._accent  = accent
        self._btn_w   = width
        self._btn_h   = height
        self._draw()
        self.bind("<Enter>",        self._on_enter)
        self.bind("<Leave>",       self._on_leave)
        self.bind("<Button-1>",    self._on_press)
        self.bind("<ButtonRelease-1>", self._on_release)

    def _draw(self, state="normal"):
        self.delete("all")
        r = 8
        if self._accent:
            if state == "hover":   bg = COLORS["accent_hover"]
            elif state == "press": bg = COLORS["accent_press"]
            else:                  bg = COLORS["accent"]
            fg = COLORS["text_dark"]
        else:
            if state == "hover":   bg = COLORS["surface_alt"]
            elif state == "press": bg = COLORS["border"]
            else:                  bg = COLORS["surface"]
            fg = COLORS["text"]

        self._round_rect(0, 0, self._btn_w, self._btn_h, r, fill=bg, outline="")
        self.create_text(self._btn_w // 2, self._btn_h // 2, text=self._text,
                         font=FONT_BODY, fill=fg, width=self._btn_w - 16)

    def _round_rect(self, x1, y1, x2, y2, r, **kw):
        self.create_polygon(
            [x1+r,y1, x2-r,y1, x2,y1+r, x2,y2-r, x2-r,y2, x1+r,y2, x1,y2-r, x1,y1+r],
            smooth=True, **kw)

    def _on_enter(self,  _e): self._draw("hover")
    def _on_leave(self,  _e): self._draw("normal")
    def _on_press(self,  _e): self._draw("press")
    def _on_release(self, _e):
        self._draw("hover")
        if self._cmd:
            self._cmd()

    def configure_state(self, disabled: bool):
        for seq in ("<Enter>","<Leave>","<Button-1>","<ButtonRelease-1>"):
            self.unbind(seq)
        if not disabled:
            self.bind("<Enter>",        self._on_enter)
            self.bind("<Leave>",       self._on_leave)
            self.bind("<Button-1>",    self._on_press)
            self.bind("<ButtonRelease-1>", self._on_release)
            self._draw("normal")
        else:
            self._draw_disabled()

    def _draw_disabled(self):
        self.delete("all")
        self._round_rect(0, 0, self._btn_w, self._btn_h, 8, fill=COLORS["border"], outline="")
        self.create_text(self._btn_w // 2, self._btn_h // 2, text=self._text,
                         font=FONT_BODY, fill=COLORS["text_dim"])


class LogPanel(tk.Frame):

    def __init__(self, master, **kw):
        super().__init__(master, bg=COLORS["log_bg"], **kw)
        scroll = tk.Scrollbar(self, orient="vertical", bd=0,
                              troughcolor=COLORS["log_bg"], activebackground=COLORS["border"])
        self.text = tk.Text(self, wrap="word", font=FONT_LOG, bd=0, padx=10, pady=8,
                            bg=COLORS["log_bg"], fg=COLORS["text"],
                            insertbackground=COLORS["text"],
                            yscrollcommand=scroll.set, state="disabled", exportselection=True)
        scroll.config(command=self.text.yview)
        self.text.grid(row=0, column=0, sticky="nsew", padx=(0,0))
        scroll.grid(row=0, column=1, sticky="ns")
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)

        self.text.tag_configure("info",    foreground=COLORS["log_info"])
        self.text.tag_configure("success", foreground=COLORS["log_success"])
        self.text.tag_configure("warning", foreground=COLORS["log_warning"])
        self.text.tag_configure("error",   foreground=COLORS["log_error"])
        self.text.tag_configure("normal",  foreground=COLORS["text"])

    def log(self, msg: str, level: str = "normal"):
        self.text.config(state="normal")
        self.text.insert("end", msg + "\n", level)
        self.text.see("end")
        self.text.config(state="disabled")

    def clear(self):
        self.text.config(state="normal")
        self.text.delete("1.0", "end")
        self.text.config(state="disabled")


class BrailleConverterApp(tk.Tk):

    WIN_W, WIN_H = 680, 720

    def __init__(self):
        super().__init__()
        self.title("Braille Converter")
        self.configure(bg=COLORS["bg"])
        self.geometry(f"{self.WIN_W}x{self.WIN_H}")
        self.minsize(self.WIN_W, 600)
        self.resizable(True, True)

        self.update_idletasks()
        sx = (self.winfo_screenwidth()  - self.WIN_W) // 2
        sy = (self.winfo_screenheight() - self.WIN_H) // 2
        self.geometry(f"+{sx}+{sy}")

        self.processor = None
        self._running  = False

        self._build_ui()

    def _build_ui(self):
        outer = tk.Frame(self, bg=COLORS["bg"])
        outer.pack(fill="both", expand=True)

        title_bar = tk.Frame(outer, bg=COLORS["bg"], pady=18)
        title_bar.pack(fill="x")
        tk.Label(title_bar, text="⠋⠗⠁⠊⠑", font=(FONT_FAMILY, 28, "bold"),
                 bg=COLORS["bg"], fg=COLORS["accent"]).pack()
        tk.Label(title_bar, text="Braille Converter  ·  PDF & Image → Braille",
                 font=(FONT_FAMILY, 11), bg=COLORS["bg"], fg=COLORS["text_dim"]).pack()

        pad_x = 20

        self._input_card(outer, pad_x)
        self._export_card(outer, pad_x)

        btn_frame = tk.Frame(outer, bg=COLORS["bg"])
        btn_frame.pack(pady=(14, 6))
        self.btn_convert = StyledButton(btn_frame, text="Convert", command=self._start_convert,
                                        accent=True, width=200, height=42)
        self.btn_convert.pack()

        self.stats_frame = tk.Frame(outer, bg=COLORS["bg"])
        self.stats_frame.pack(fill="x", padx=pad_x, pady=(4, 2))

        log_wrapper = tk.LabelFrame(outer, text=" Log ", font=(FONT_FAMILY, 10, "bold"),
                                    bg=COLORS["surface"], fg=COLORS["text_dim"],
                                    bd=1, relief="groove", highlightcolor=COLORS["border"])
        log_wrapper.pack(fill="both", expand=True, padx=pad_x, pady=(2, 16))
        log_wrapper.grid_rowconfigure(0, weight=1)
        log_wrapper.grid_columnconfigure(0, weight=1)

        self.log_panel = LogPanel(log_wrapper)
        self.log_panel.grid(row=0, column=0, sticky="nsew", padx=4, pady=4)

        self.log("Welcome!  Browse to a PDF or image file and press Convert.", "info")

    def _input_card(self, parent, pad_x):
        card = tk.LabelFrame(parent, text=" Input & Output ", font=(FONT_FAMILY, 10, "bold"),
                             bg=COLORS["surface"], fg=COLORS["text_dim"],
                             bd=1, relief="groove", highlightcolor=COLORS["border"])
        card.pack(fill="x", padx=pad_x, pady=(0, 6))
        card.grid_columnconfigure(1, weight=1)

        tk.Label(card, text="Source file", font=FONT_BODY,
                 bg=COLORS["surface"], fg=COLORS["text"]).grid(
                     row=0, column=0, padx=(12, 6), pady=(10, 4), sticky="w")

        self.input_var = tk.StringVar()
        entry_in = tk.Entry(card, textvariable=self.input_var, font=FONT_BODY,
                            bd=0, bg=COLORS["log_bg"], fg=COLORS["text"],
                            insertbackground=COLORS["text"], relief="flat",
                            highlightthickness=1, highlightcolor=COLORS["accent"],
                            highlightbackground=COLORS["border"])
        entry_in.grid(row=0, column=1, pady=(10, 4), ipady=6, sticky="ew", padx=(0, 6))
        self.input_entry = entry_in

        browse_in = StyledButton(card, text="Browse …", command=self._browse_input,
                                 accent=False, width=96, height=32)
        browse_in.grid(row=0, column=2, padx=(0, 12), pady=(10, 4))

        tk.Label(card, text="Output folder", font=FONT_BODY,
                 bg=COLORS["surface"], fg=COLORS["text"]).grid(
                     row=1, column=0, padx=(12, 6), pady=(4, 10), sticky="w")

        self.output_var = tk.StringVar(value=os.getcwd())
        entry_out = tk.Entry(card, textvariable=self.output_var, font=FONT_BODY,
                             bd=0, bg=COLORS["log_bg"], fg=COLORS["text"],
                             insertbackground=COLORS["text"], relief="flat",
                             highlightthickness=1, highlightcolor=COLORS["accent"],
                             highlightbackground=COLORS["border"])
        entry_out.grid(row=1, column=1, pady=(4, 10), ipady=6, sticky="ew", padx=(0, 6))

        browse_out = StyledButton(card, text="Browse …", command=self._browse_output,
                                  accent=False, width=96, height=32)
        browse_out.grid(row=1, column=2, padx=(0, 12), pady=(4, 10))

    def _export_card(self, parent, pad_x):
        card = tk.LabelFrame(parent, text=" Export Options ", font=(FONT_FAMILY, 10, "bold"),
                             bg=COLORS["surface"], fg=COLORS["text_dim"],
                             bd=1, relief="groove", highlightcolor=COLORS["border"])
        card.pack(fill="x", padx=pad_x, pady=(0, 0))

        self.chk_txt     = tk.BooleanVar(value=True)
        self.chk_normal  = tk.BooleanVar(value=True)
        self.chk_braille = tk.BooleanVar(value=True)

        for col, (var, label, tip) in enumerate([
            (self.chk_txt,     "Braille TXT",  "Plain-text file with Braille Unicode"),
            (self.chk_normal,  "Normal DOCX",  "Word doc with original text (RTL for Arabic)"),
            (self.chk_braille, "Braille DOCX", "Word doc containing Braille text"),
        ]):
            cb = tk.Checkbutton(card, variable=var, text=label, font=FONT_BODY,
                                bg=COLORS["surface"], fg=COLORS["text"],
                                activebackground=COLORS["surface"],
                                activeforeground=COLORS["text"],
                                selectcolor=COLORS["log_bg"],
                                indicatoron=True, bd=0)
            cb.grid(row=0, column=col, padx=(12 if col == 0 else 24, 0), pady=10, sticky="w")
            ToolTip(cb, tip)

    def _browse_input(self):
        path = filedialog.askopenfilename(
            title="Select input file",
            filetypes=[("Supported", "*.pdf *.png *.jpg *.jpeg"),
                       ("PDF", "*.pdf"),
                       ("Images", "*.png *.jpg *.jpeg"),
                       ("All", "*.*")]
        )
        if path:
            self.input_var.set(path)

    def _browse_output(self):
        folder = filedialog.askdirectory(title="Select output folder")
        if folder:
            self.output_var.set(folder)

    def log(self, msg: str, level: str = "normal"):
        self.log_panel.log(msg, level)

    def _clear_log(self):
        self.log_panel.clear()

    def _clear_stats(self):
        for w in self.stats_frame.winfo_children():
            w.destroy()

    def _show_stats(self, stats: dict):
        self._clear_stats()
        pills = [
            ("Language",  stats["primary_language"].upper(), COLORS["accent"]),
            ("Arabic",    str(stats["arabic_chars"]),        COLORS["warning"]),
            ("English",   str(stats["english_chars"]),       COLORS["log_info"]),
            ("Total",     str(stats["total_chars"]),         COLORS["text"]),
        ]
        for i, (label, value, colour) in enumerate(pills):
            pill = tk.Frame(self.stats_frame, bg=COLORS["surface"], bd=0)
            pill.pack(side="left", padx=(0 if i == 0 else 6, 0))
            tk.Label(pill, text=label, font=(FONT_FAMILY, 9),
                     bg=COLORS["surface"], fg=COLORS["text_dim"]).pack(side="left", padx=(8, 2))
            tk.Label(pill, text=value, font=FONT_STAT,
                     bg=COLORS["surface"], fg=colour).pack(side="left", padx=(0, 8))

    def _start_convert(self):
        if self._running:
            return
        input_path = self.input_var.get().strip()
        if not input_path:
            messagebox.showwarning("No file", "Please select an input file first.")
            return
        if not os.path.isfile(input_path):
            messagebox.showerror("File not found", f"Cannot find:\n{input_path}")
            return

        self._running = True
        self.btn_convert.configure_state(disabled=True)
        self._clear_log()
        self._clear_stats()

        t = threading.Thread(target=self._do_convert, args=(input_path,), daemon=True)
        t.start()

    def _do_convert(self, input_path: str):
        try:
            self.processor = FileProcessor(log_callback=self._thread_log)

            self._thread_log("─" * 52, "info")
            self._thread_log("  Braille Converter  ·  Starting …", "info")
            self._thread_log("─" * 52, "info")

            text = self.processor.process_file(input_path)
            self._thread_log(f"✓  Extracted {len(text)} characters.", "success")

            stats = self.processor.converter.get_language_stats(text)
            self.after(0, self._show_stats, stats)

            self._thread_log("Converting to Braille …", "info")
            braille = self.processor.converter.text_to_braille(text)
            self._thread_log("✓  Braille conversion complete.", "success")

            out_dir = self.output_var.get().strip() or os.getcwd()
            os.makedirs(out_dir, exist_ok=True)
            base = Path(input_path).stem

            saved_any = False

            if self.chk_txt.get():
                txt_path = os.path.join(out_dir, f"{base}_braille.txt")
                with open(txt_path, 'w', encoding='utf-8') as fh:
                    fh.write(braille)
                self._thread_log(f"✓  TXT  →  {txt_path}", "success")
                saved_any = True

            if self.chk_normal.get():
                try:
                    docx_path = os.path.join(out_dir, f"{base}_normal.docx")
                    self.processor.save_normal_docx(text, docx_path)
                    self._thread_log(f"✓  Normal DOCX  →  {docx_path}", "success")
                    saved_any = True
                except ImportError as e:
                    self._thread_log(f"⚠  Normal DOCX skipped: {e}", "warning")
                except Exception as e:
                    self._thread_log(f"✗  Normal DOCX error: {e}", "error")

            if self.chk_braille.get():
                try:
                    docx_path = os.path.join(out_dir, f"{base}_braille.docx")
                    self.processor.save_braille_docx(braille, docx_path)
                    self._thread_log(f"✓  Braille DOCX  →  {docx_path}", "success")
                    saved_any = True
                except ImportError as e:
                    self._thread_log(f"⚠  Braille DOCX skipped: {e}", "warning")
                except Exception as e:
                    self._thread_log(f"✗  Braille DOCX error: {e}", "error")

            self._thread_log("─" * 52, "info")
            if saved_any:
                self._thread_log("  All done!  Check your output folder.", "success")
            else:
                self._thread_log("  No files were saved — enable at least one export option.", "warning")
            self._thread_log("─" * 52, "info")

        except Exception as exc:
            self._thread_log(f"✗  {exc}", "error")
            self._thread_log(traceback.format_exc(), "error")

        finally:
            self.after(0, self.btn_convert.configure_state, False)
            self.after(0, setattr, self, "_running", False)

    def _thread_log(self, msg: str, level: str = "normal"):
        self.after(0, self.log, msg, level)


def main():
    app = BrailleConverterApp()
    app.mainloop()


if __name__ == "__main__":
    main()
